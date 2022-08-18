# #!/usr/bin/python

import sys
import textProcessor
import docx
import copy


from PyQt6.QtWidgets import (
    QApplication,
    QGridLayout,
    QPushButton,
    QLabel,
    QProgressBar,
    QWidget,
    QLineEdit,
    QFileDialog,
    QMainWindow,

)
from PyQt6 import QtGui
from PyQt6.QtCore import (
    QThread,
    QObject,
    pyqtSignal
)

import os


basedir = os.path.dirname(__file__)

try:
    from ctypes import windll  # Only exists on Windows.
    myappid = 'rusubtitles.com/'
    windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
except ImportError:
    pass


class ProcessDocument(QObject):
    progress = pyqtSignal(int)
    started = pyqtSignal()
    finished = pyqtSignal()

    def __init__(self, fname_txt, fname_docx, final_file_name):
        super().__init__()
        self.fname_txt = fname_txt
        self.fname_docx = fname_docx
        self.final_file_name = final_file_name

    def run(self):
        #("Thread start")
        self.started.emit()

        self.process_files()
        #print("Thread complete")
        self.finished.emit()

    def process_files(self):

        doc = docx.Document(self.fname_docx)

        table_docx = doc.tables[0]

        self.progress.emit(0)
        characters_with_colors = textProcessor.get_characters_add_colors(self.fname_txt)
        #(characters_with_colors)
        self.progress.emit(10)

        textProcessor.set_color_for_characters(characters_with_colors, table_docx, characters_column=1)
        self.progress.emit(25)
        textProcessor.set_colors_for_text(table_docx, text_column=2)
        self.progress.emit(50)

        # create new dict with characters and symbols number
        characters_words_zero = copy.deepcopy(characters_with_colors)
        for actor, characters in characters_words_zero.items():
            for character, color in characters.items():
                characters[character] = 0
        # ---
        self.progress.emit(55)

        textProcessor.check_time_code_parameters(table_docx)
        self.progress.emit(75)

        characters_with_number_words = textProcessor.count_character_words(characters_words_zero, table_docx)
        #print(characters_with_number_words)
        textProcessor.docx_add_counted_characters(doc, table_docx, characters_with_number_words, characters_with_colors)
        self.progress.emit(90)

        textProcessor.set_size_table_border(table_docx, 4, 'black')
        self.progress.emit(95)

        doc.save(self.final_file_name)
        self.progress.emit(100)


class Window(QMainWindow):
    def __init__(self):
        super().__init__()

        self.fname_txt = None
        self.fname_docx = None
        self.final_file_name = ""

        abspath = os.path.abspath("__file__")
        self.this_path_name = os.path.dirname(abspath)

        self.setWindowTitle("RuFilms")

        # Create a QGridLayout instance
        self.layout = QGridLayout()

        # Create labels
        self.label_txt = QLabel('Choose .txt document (actors, characters, colors)', self)
        self.label_txt.setStyleSheet("background-color: white; border: 1px solid black;")
        self.label_txt.setFixedSize(300, 20)

        self.label_docx = QLabel('Choose .docx document (table)', self)
        self.label_docx.setStyleSheet("background-color: white; border: 1px solid black;")
        self.label_docx.setFixedSize(300, 20)

        self.information_label = QLabel('Information: choose files', self)

        # Create buttons
        self.choose_txt_button = QPushButton("Choose characters", self)
        self.choose_txt_button.clicked.connect(self.choose_txt_file)

        self.choose_docx_button = QPushButton("Choose docx", self)
        self.choose_docx_button.clicked.connect(self.choose_docx_file)

        self.start_processing_button = QPushButton("Start")
        self.start_processing_button.clicked.connect(self.start_process)
        self.start_processing_button.setEnabled(False)

        # Create text editor
        self.final_file_name_text_editor = QLineEdit("Add final file name")
        self.final_file_name_text_editor.textChanged[str].connect(self.check_status_for_start)
        self.final_file_name_label = QLabel("Final file name")

        # Create progress bar
        self.prog_bar_file_processing = QProgressBar(self)
        self.prog_bar_file_processing.setValue(0)

        # Add widgets to the layout

        self.layout.addWidget(self.information_label, 0, 0, 1, 2)

        self.layout.addWidget(self.label_txt, 1, 0)
        self.layout.addWidget(self.choose_txt_button, 1, 1)

        self.layout.addWidget(self.label_docx, 2, 0)
        self.layout.addWidget(self.choose_docx_button, 2, 1)

        self.layout.addWidget(self.final_file_name_text_editor, 3, 0, 1, 1)
        self.layout.addWidget(self.final_file_name_label, 3, 1)

        self.layout.addWidget(self.prog_bar_file_processing, 4, 0)
        self.layout.addWidget(self.start_processing_button, 4, 1)

        self.layout.setVerticalSpacing(20)
        # Set the layout on the application's window
        self.widget = QWidget()
        self.widget.setLayout(self.layout)
        self.setCentralWidget(self.widget)

    def choose_txt_file(self):
        self.fname_txt = QFileDialog.getOpenFileName(self, 'Open file', self.this_path_name, "txt(*.txt)")[0]
        if self.fname_txt:
            self.label_txt.setText(self.fname_txt)
            self.check_status_for_start(None)

    def choose_docx_file(self):
        self.fname_docx = QFileDialog.getOpenFileName(self, 'Open file', self.this_path_name, "docx(*.docx)")[0]
        if self.fname_docx:
            self.label_docx.setText(self.fname_docx)
            self.final_file_name = self.fname_docx[:-5] + "_Final"
            self.final_file_name_text_editor.setText(os.path.basename(self.final_file_name))
            self.check_status_for_start(None)

    def check_status_for_start(self, final_file_name):

        if final_file_name:
            final_file_name += ".docx"
            all_dir_files = [f for f in os.listdir(self.this_path_name) if
                             os.path.isfile(os.path.join(self.this_path_name, f))]
            for file in all_dir_files:
                if final_file_name == file:
                    #print("bad_status")
                    self.information_label.setText("Information: a file with the same name already exists!")
                    self.information_label.setStyleSheet("QLabel {color : red}")
                    self.start_processing_button.setEnabled(False)
                    return None

            if (self.fname_docx != None) and (self.fname_txt != None):
                #print("good_status")
                self.final_file_name = final_file_name
                self.information_label.setText('Information: press "Start" button.')
                self.information_label.setStyleSheet("QLabel {color : green}")
                self.start_processing_button.setEnabled(True)
                return None

            if (self.fname_docx != None) or (self.fname_txt != None):
                #("middle_status")
                self.information_label.setText('Information: choose files')
                self.information_label.setStyleSheet("QLabel {color : black}")
                self.start_processing_button.setEnabled(False)

    def start_process(self):

        # Step 2: Create a QThread object
        self.thread = QThread()
        # Step 3: Create a worker object
        self.worker = ProcessDocument(self.fname_txt, self.fname_docx, self.final_file_name)
        self.worker.moveToThread(self.thread)

        self.thread.started.connect(self.worker.run)
        self.worker.started.connect(self.process_thread_start)
        self.worker.finished.connect(self.process_thread_finished)
        self.worker.progress.connect(self.reportProgress)

        self.thread.start()


    def process_thread_start(self):
        self.information_label.setText('Information: files are processed')
        self.information_label.setStyleSheet("QLabel {color : orange}")

        self.start_processing_button.setEnabled(False)
        self.choose_docx_button.setEnabled(False)
        self.choose_txt_button.setEnabled(False)
        self.final_file_name_text_editor.setEnabled(False)

    def reportProgress(self, value_progress_bar):
        self.prog_bar_file_processing.setValue(value_progress_bar)

    def process_thread_finished(self):
        self.information_label.setText('Information: all done, choose another files')
        self.information_label.setStyleSheet("QLabel {color : black}")
        self.start_processing_button.setEnabled(True)
        self.choose_docx_button.setEnabled(True)
        self.choose_txt_button.setEnabled(True)
        self.final_file_name_text_editor.setEnabled(True)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = Window()
    app.setWindowIcon(QtGui.QIcon('D:\\Progects\\rufilms\\RuFilms.ico'))
    window.show()
    sys.exit(app.exec())
