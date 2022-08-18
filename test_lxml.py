import re

import docx
import styles
import textProcessor


def set_size_table_border(table_docx, border_size, color):
    tbl = table_docx._tbl
    tblPr = tbl.tblPr

    tblBorders = docx.oxml.OxmlElement('w:tblBorders')

    tbl_borders = docx.oxml.OxmlElement('w:tblBorders')
    w_names = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    w_params = {'val': 'single', 'sz': str(border_size), 'space': "0", 'color': color}
    for name in w_names:
        w_name = docx.oxml.OxmlElement('w:' + name)
        for param, value in w_params.items():
            w_name.set(docx.oxml.ns.qn('w:' + param), value)
        tblBorders.append(w_name)

    tblPr.append(tblBorders)
    return tblPr


def check_time_code_parameters(table_docx):
    re_minutes_seconds_colon = re.compile("^([0-5]?[0-9]):([0-5]?[0-9])$")
    # re_ours_minutes_seconds_colon = re.compile("^(2[0-3]|[01]?[0-9]):([0-5]?[0-9]):([0-5]?[0-9])$")
    current_time_code = None
    previous_time_code = None

    for i, row in enumerate(table_docx.rows):
        paragraphs = row.cells[0].paragraphs
        for paragraph in paragraphs:
            text = paragraph.text

            if (len(text) > 0) and (text != "\n"):
                match_result = re.match(re_minutes_seconds_colon, text)
                if bool(match_result):
                    minutes, seconds = match_result.group(0).split(':', 1)
                    current_time_code = int(minutes) * 60 + int(seconds)
                    if previous_time_code:
                        difference = current_time_code - previous_time_code
                        # wrong time
                        if difference <= 0:
                            for run in paragraph.runs:
                                styles.text_highlight_red(run)
                # wrong text format
                else:
                    for run in paragraph.runs:
                        styles.text_highlight_red(run)
                previous_time_code = current_time_code


if __name__ == '__main__':
    doc_name = "New_Life_1_2"

    doc = docx.Document(doc_name + '.docx')

    table_docx = doc.tables[0]

    characters_with_colors = textProcessor.get_characters_add_colors('Characters_new_life_1_2.txt')
