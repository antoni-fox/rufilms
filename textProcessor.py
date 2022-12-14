import docx
import styles
import re


def get_characters_add_colors(file_name):
    characters = {}
    characters_colors = {}
    with open(file_name, encoding='utf8') as file:
        lines = file.readlines()
        color_number = 0
        for line in lines:
            if (line and not line.isspace()):
                data = line.split(':', 2)
                actor_name = data[0].strip()
                characters_names = data[1].split(',')
                data[2] = data[2].strip()
                color = [int(number) for number in data[2][data[2].find("(") + 1: data[2].find(")")].split(",")]
                characters[actor_name] = {}

                for name in characters_names:
                    characters[actor_name][name.strip().upper()] = color
                    # characters_colors[name] = styles.COLORS[color_number]

    return characters


def read_table(table):
    table_data = []
    keys = None

    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        # Establish the mapping based on the first row
        # headers; these will become the keys of our dictionary
        if i == 0:
            keys = tuple(text)
            continue

        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        table_data.append(row_data)

    return table_data


def set_color_for_characters(characters, table_docx, characters_column):
    all_names_colors = {}
    for actor, key in characters.items():
        all_names_colors.update(key)
    # print(all_names_colors)
    for i, row in enumerate(table_docx.rows):
        cell_xml_element = row.cells[characters_column]
        for idx, paragraph in enumerate(cell_xml_element.paragraphs):
            text_paragraph = paragraph.text.strip().upper()

            dict_position_name = {}
            for name, color in all_names_colors.items():
                result = text_paragraph.find(name)
                if result != (-1):
                    paragraph = paragraph.clear()
                    dict_position_name[result] = name
            if bool(dict_position_name):
                keys = list(dict_position_name.keys())
                keys.sort()
                for number in keys:
                    run = paragraph.add_run(dict_position_name[number] + " ")
                    styles.characters_style(run, all_names_colors[dict_position_name[number]])


def set_colors_for_text(table_docx, text_column):
    # regexp patterns for searching words
    reg_pattern_sounds = re.compile(r'\*[^*]*\*')
    reg_pattern_voiceover = re.compile(r'\(??/??\)|\(??/??\)')
    # ----

    # iterate from rows in table
    for i, row in enumerate(table_docx.rows):
        cell_xml_element = row.cells[text_column]

        # iterate from paragraphs cell table
        for paragraph in cell_xml_element.paragraphs:
            text = paragraph.text

            # Search all bold worlds
            all_bold_words = {}
            start_of_run = 0
            end_of_run = 0
            for run in paragraph.runs:
                end_of_run += len(run.text)
                if run.bold:
                    all_bold_words[start_of_run] = [run.text, styles.standart_text_bold]
                start_of_run = end_of_run
            # ----

            paragraph = paragraph.clear()
            all_words_sounds = {}
            all_words_voiceover = {}

            # search all words and start position of word
            for m in reg_pattern_sounds.finditer(text):
                all_words_sounds[m.start()] = [m.group(), styles.actor_sounds_text]

            for m in reg_pattern_voiceover.finditer(text):
                all_words_voiceover[m.start()] = [m.group(), styles.actor_voiceover]
            # ----

            # concatenate all dict and sort text by sorting in ascending order start index
            mixed_words = {**all_words_sounds, **all_words_voiceover, **all_bold_words}
            keys_mixed_words = list(mixed_words.keys())
            keys_mixed_words.sort()
            # ----

            # add all styled text in paragraph
            previous_word_position = 0
            for start_word_position in keys_mixed_words:
                word = str(mixed_words[start_word_position][0])
                end_word_position = start_word_position + len(word)

                run_start = paragraph.add_run(text[previous_word_position:start_word_position])
                styles.standart_text(run_start)

                run_word = paragraph.add_run(word)
                mixed_words[start_word_position][1](run_word)

                previous_word_position = end_word_position

            run_end = paragraph.add_run(text[previous_word_position:])
            styles.standart_text(run_end)
            # ---


def count_character_words(characters, table_docx):
    for i, row in enumerate(table_docx.rows):
        characters_text = row.cells[1].text
        voice_text = row.cells[2].text

        # print(voice_text)

        for actor, characters_dict in characters.items():
            for character_name, symbols in characters_dict.items():
                text_without_symbols = re.sub(r'\(??/??\)|\(??/??\)|/|//|\.\.|\.\.\.', '', voice_text)
                # print(text_without_symbols)
                words_number = len(re.findall(r'\w+', text_without_symbols))
                # print(words_number)
                result = characters_text.find(character_name)
                if result != (-1):
                    characters[actor][character_name] = symbols + words_number

    return characters


def docx_add_counted_characters(doc, table_docx, characters_with_number_words, characters_with_colors):
    for actor, characters in characters_with_number_words.items():
        new_paragraph = doc.add_paragraph("")
        paragraph_run = new_paragraph.add_run(actor + " - ")
        styles.characters_style(paragraph_run, list(characters_with_colors[actor].values())[0])
        for character, number_words in characters.items():
            if number_words > 0:
                run = new_paragraph.add_run(character.capitalize() + " ")
                styles.characters_style(run, characters_with_colors[actor][character])
                run = new_paragraph.add_run(str(number_words) + ", ")
                styles.characters_style(run, [0, 0, 0])
                table_docx._element.addprevious(new_paragraph._p)
        new_paragraph = doc.add_paragraph("")
        paragraph_run = new_paragraph.add_run(str(sum(list(characters.values()))))
        styles.standart_text(paragraph_run)
        table_docx._element.addprevious(new_paragraph._p)


def set_size_table_border(table_docx, border_size, color):
    tbl = table_docx._tbl
    tblPr = tbl.tblPr

    tblBorders = docx.oxml.OxmlElement('w:tblBorders')

    tbl_borders = docx.oxml.OxmlElement('w:tblBorders')
    w_names = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    w_params = {'val': 'single', 'sz': str(border_size), 'space': "0", 'color': color}
    for name in w_names:
        w_name = docx.oxml.OxmlElement('w:' + name)
        for param, value in w_params.items():
            w_name.set(docx.oxml.ns.qn('w:' + param), value)
        tblBorders.append(w_name)

    tblPr.append(tblBorders)
    return tblPr


def check_time_code_parameters(table_docx):
    re_minutes_seconds_colon = re.compile("^([0-5]?[0-9]):([0-5]?[0-9])$")
    # re_ours_minutes_seconds_colon = re.compile("^(2[0-3]|[01]?[0-9]):([0-5]?[0-9]):([0-5]?[0-9])$")
    current_time_code = None
    previous_time_code = None

    for i, row in enumerate(table_docx.rows):
        paragraphs = row.cells[0].paragraphs
        for paragraph in paragraphs:
            text = paragraph.text

            if (len(text) > 0) and (text != "\n"):
                match_result = re.match(re_minutes_seconds_colon, text)
                if bool(match_result):
                    minutes, seconds = match_result.group(0).split(':', 1)
                    current_time_code = int(minutes) * 60 + int(seconds)
                    if previous_time_code:
                        difference = current_time_code - previous_time_code
                        # wrong time
                        if difference < 0:
                            for run in paragraph.runs:
                                styles.text_highlight_red(run)
                # wrong text format
                else:
                    for run in paragraph.runs:
                        styles.text_highlight_red(run)
                previous_time_code = current_time_code


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
